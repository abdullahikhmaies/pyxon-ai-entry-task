from docx import Document

doc = Document()

# Add a Title
doc.add_heading('مقدمة في الذكاء الاصطناعي', 0)

# Add a Paragraph with diacritics
p = doc.add_paragraph('الذَّكَاءُ الاصْطِنَاعِيُّ هو مَجَالٌ يَهْتَمُّ بِبِنَاءِ أَنْظِمَةٍ قَادِرَةٍ عَلَى القِيَامِ بِمَهَامٍ تَتَطَلَّبُ ذَكَاءً بَشَرِيًّا.')

# Add a Chapter Heading (to trigger dynamic chunking)
doc.add_heading('الفصل الأول: تَعَلُّمُ الآلَةِ', level=1)

p2 = doc.add_paragraph('يَعْتَمِدُ تَعَلُّمُ الآلَةِ عَلَى خَوَارِزْمِيَّاتٍ تَتَعَلَّمُ مِنَ البَيَانَاتِ بَدَلًا مِنْ بَرْمَجَتِهَا بِشَكْلٍ صَرِيحٍ.')
p2 = doc.add_paragraph('هناك أنواع متعددة مثل التعلم بالإشراف والتعلم بدون إشراف.')

doc.add_heading('خاتمة', level=1)
doc.add_paragraph('نأمل أن يكون هذا العرض التوضيحي قد نال إعجابكم.')

doc.save('data/sample_demo.docx')
print("Created data/sample_demo.docx")
